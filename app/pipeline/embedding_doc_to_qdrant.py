import os
import docx
from docx.text.paragraph import Paragraph
from docx.table import Table
from sentence_transformers import SentenceTransformer
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams, PointStruct
import logging
import uuid

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

MODEL_NAME = 'BAAI/bge-base-en-v1.5'
QDRANT_HOST = "qdrant_db"  
QDRANT_PORT = 6333
VECTOR_SIZE = 768
COLLECTION_NAME = "policy_and_size"

# chuyển table trong word thành markdown table
def table_to_markdown(table):
    markdown_lines = []
    rows = list(table.rows)
    if not rows:
        return ""
    
    header_cells = [cell.text.strip().replace("\n", " ") for cell in rows[0].cells]
    markdown_lines.append("| " + " | ".join(header_cells) + " |")
    markdown_lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
    
    for row in rows[1:]:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        markdown_lines.append("| " + " | ".join(cells) + " |")
        
    return "\n".join(markdown_lines)

# duyệt tuần tự để lấy cả paragraph và table theo thứ tự hiển thị 
def parse_elements_sequences(doc):
    elements = []
    for child in doc.element.body.iterchildren():
        if child.tag.endswith('p'):
            p = Paragraph(child, doc)
            if p.text.strip():
                elements.append(("text", p.text.strip()))
        elif child.tag.endswith('tbl'):
            t = Table(child, doc)
            md_table = table_to_markdown(t)
            if md_table:
                elements.append(("table", md_table))
    return elements

# structure-based + cumulative character chunking: gom gạch đầu dòng theo heading
def chunk_return_policy(file_path, max_chunk_size=600):
    if not os.path.exists(file_path):
        logger.error(f"Không tìm thấy file: {file_path}")
        return []
    
    doc = docx.Document(file_path)
    elements = parse_elements_sequences(doc)
    chunks = []
    
    known_headings = [
        "Return and Exchange Deadline",
        "Items required for Returns and Exchanges",
        "Products that Cannot be Exchanged or Returned",
        "Additional items required for Exchange or Return"
    ]
    
    current_context = "Return and Exchange Policy"
    current_pieces = []
    current_len = 0
    
    for el_type, el_content in elements:
        if el_type == "text" and el_content in known_headings:
            if current_pieces:
                full_content = "\n".join(current_pieces)
                chunks.append({
                    "context": current_context,
                    "type": "text",
                    "full_text": f"[{current_context}]\n{full_content}"
                })
                current_pieces = []
                current_len = 0
            current_context = el_content
            continue
        
        current_pieces.append(el_content)
        current_len += len(el_content)
        
        if current_len >= max_chunk_size:
            full_content = "\n".join(current_pieces)
            chunks.append({
                "context": current_context,
                "type": "text",
                "full_text": f"[{current_context}]\n{full_content}"
            })
            current_pieces = []
            current_len = 0
            
    if current_pieces:
        full_content = "\n".join(current_pieces)
        chunks.append({
            "context": current_context,
            "type": "text",
            "full_text": f"[{current_context}]\n{full_content}"
        })
        
    return chunks

# structure-based chunking: tách đoạn văn, giữ nguyên các bảng
def chunk_delivery_policy(file_path):
    if not os.path.exists(file_path):
        logger.error(f"Not found: {file_path}")
        return []
        
    doc = docx.Document(file_path)
    elements = parse_elements_sequences(doc)
    chunks = []
    
    current_context = "Delivery Policy"
    
    for el_type, el_content in elements:
        if el_type == "text" and "Home Delivery" in el_content:
            current_context = "Home Delivery"
            
        if el_type == "table":
            chunks.append({
                "context": current_context,
                "type": "table",
                "full_text": f"[{current_context} Table]\n{el_content}"
            })
        else:
            chunks.append({
                "context": current_context,
                "type": "text",
                "full_text": f"[{current_context}] {el_content}"
            })
    return chunks

# table-level / whole-file chunking: gom trọn vẹn bảng size kèm quy tắc chọn size vào chung một chunk
def chunk_size_guide(file_path):
    if not os.path.exists(file_path):
        logger.error(f"Not found: {file_path}")
        return []
        
    doc = docx.Document(file_path)
    elements = parse_elements_sequences(doc)
    
    # quy tắc nếu không bóc tách được text
    size_rule = "Size selection rule: If the height and weight fall into two different sizes, choose the larger size." 
    for el_type, el_content in elements:
        if "Size selection rule" in el_content:
            size_rule = el_content
            break
            
    chunks = []
    for el_type, el_content in elements:
        if "Men Size" in el_content or (el_type == "table" and "Men Size" in elements[elements.index((el_type, el_content))-1][1]):
            if el_type == "table":
                chunks.append({
                    "context": "Men Size Guide",
                    "type": "table",
                    "full_text": f"[Men Size Table]\n{el_content}\n\n*Important Rule: {size_rule}"
                })
        elif "Women Size" in el_content or (el_type == "table" and "Women Size" in elements[elements.index((el_type, el_content))-1][1]):
            if el_type == "table":
                chunks.append({
                    "context": "Women Size Guide",
                    "type": "table",
                    "full_text": f"[Women Size Table]\n{el_content}\n\n*Important Rule: {size_rule}"
                })
    return chunks

def main():
    logger.info(f"Loading Embedding Model...")
    model = SentenceTransformer(MODEL_NAME)
    
    qdrant_client = QdrantClient(host=QDRANT_HOST, port=QDRANT_PORT)
    
    try:
        qdrant_client.get_collection(collection_name=COLLECTION_NAME)
        logger.info(f"Collection '{COLLECTION_NAME}' already exist.")
    except Exception:
        logger.info(f"Đang tạo mới kho chứa vector: '{COLLECTION_NAME}'...")
        qdrant_client.create_collection(
            collection_name=COLLECTION_NAME,
            vectors_config=VectorParams(size=VECTOR_SIZE, distance=Distance.COSINE)
        )
        
    # dường dẫn thư mục dữ liệu bên trong container 
    data_dir = "/app/data"
    
    all_chunks = []
    
    # file chính sách đổi trả
    ret_chunks = chunk_return_policy(os.path.join(data_dir, "Return and Exchange Conditions.docx"))
    for c in ret_chunks: c["source_file"] = "Return and Exchange Conditions.docx"
    all_chunks.extend(ret_chunks)
    
    # file vận chuyển 
    del_chunks = chunk_delivery_policy(os.path.join(data_dir, "Delivery Method and Cost.docx"))
    for c in del_chunks: c["source_file"] = "Delivery Method and Cost.docx"
    all_chunks.extend(del_chunks)
    
    # file size
    size_chunks = chunk_size_guide(os.path.join(data_dir, "Size.docx"))
    for c in size_chunks: c["source_file"] = "Size.docx"
    all_chunks.extend(size_chunks)
    
    if not all_chunks:
        logger.warning("No valid chunks extracted. Terminating process.")
        return
        
    logger.info(f"Total chunks created after splitting: {len(all_chunks)}")
    
    points = []
    for idx, chunk in enumerate(all_chunks):
        vector = model.encode(chunk["full_text"]).tolist()
        point = PointStruct(
            id=str(uuid.uuid4()),  # UUID string to guarantee zero ID collisions with product data
            vector=vector,
            payload={
                "source_file": chunk["source_file"],
                "context": chunk["context"],
                "type": chunk["type"],
                "full_text": chunk["full_text"]
            }
        )
        points.append(point)
        
    logger.info("Upserting structured data to Qdrant...")
    qdrant_client.upsert(collection_name=COLLECTION_NAME, points=points)
    logger.info("DONE!")

if __name__ == "__main__":
    main()